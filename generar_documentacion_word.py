from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

# Crear el directorio de documentación si no existe
doc_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'public', 'documentacion')
os.makedirs(doc_dir, exist_ok=True)

# Crear un nuevo documento
doc = Document()

# Configurar estilos
styles = doc.styles

# Estilo para título principal
style_title = styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
style_title.font.name = 'Arial'
style_title.font.size = Pt(24)
style_title.font.bold = True
style_title.font.color.rgb = RGBColor(67, 56, 202)  # Indigo-600

# Estilo para subtítulos
style_heading1 = styles.add_style('Titulo1', WD_STYLE_TYPE.PARAGRAPH)
style_heading1.font.name = 'Arial'
style_heading1.font.size = Pt(18)
style_heading1.font.bold = True
style_heading1.font.color.rgb = RGBColor(79, 70, 229)  # Indigo-500

# Estilo para subtítulos secundarios
style_heading2 = styles.add_style('Titulo2', WD_STYLE_TYPE.PARAGRAPH)
style_heading2.font.name = 'Arial'
style_heading2.font.size = Pt(16)
style_heading2.font.bold = True
style_heading2.font.color.rgb = RGBColor(99, 102, 241)  # Indigo-400

# Estilo para subtítulos terciarios
style_heading3 = styles.add_style('Titulo3', WD_STYLE_TYPE.PARAGRAPH)
style_heading3.font.name = 'Arial'
style_heading3.font.size = Pt(14)
style_heading3.font.bold = True
style_heading3.font.color.rgb = RGBColor(129, 140, 248)  # Indigo-300

# Estilo para texto normal
style_normal = styles.add_style('TextoNormal', WD_STYLE_TYPE.PARAGRAPH)
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(11)
style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
style_normal.paragraph_format.space_after = Pt(6)

# Estilo para texto de código
style_code = styles.add_style('TextoCodigo', WD_STYLE_TYPE.PARAGRAPH)
style_code.font.name = 'Courier New'
style_code.font.size = Pt(10)
style_code.paragraph_format.space_before = Pt(6)
style_code.paragraph_format.space_after = Pt(6)

# Estilo para pie de página
style_footer = styles.add_style('PiePagina', WD_STYLE_TYPE.PARAGRAPH)
style_footer.font.name = 'Arial'
style_footer.font.size = Pt(8)
style_footer.font.italic = True
style_footer.font.color.rgb = RGBColor(107, 114, 128)  # gray-500

# ----- PORTADA -----
# Agregar imagen de logo (comentado porque no hay imagen disponible)
# doc.add_picture('path/to/logo.png', width=Inches(2.5))

# Título principal
title = doc.add_paragraph('DOCUMENTACIÓN TÉCNICA Y FUNCIONAL', style='TituloPrincipal')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtítulo
subtitle = doc.add_paragraph('Sistema NovelasApp', style='Titulo1')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Versión
version = doc.add_paragraph('Versión 1.0', style='Titulo2')
version.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Fecha de generación
current_date = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
date_paragraph = doc.add_paragraph(f'Documento generado el: {current_date}', style='TextoNormal')
date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Salto de página después de la portada
doc.add_page_break()

# ----- ÍNDICE DE CONTENIDOS -----
toc_title = doc.add_paragraph('ÍNDICE DE CONTENIDOS', style='TituloPrincipal')
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Aquí iría el código para insertar un índice automático
# No es posible insertar un TOC programáticamente con python-docx de manera sencilla
# pero se puede crear uno manualmente

doc.add_paragraph('1. INTRODUCCIÓN', style='Titulo1')
doc.add_paragraph('    1.1 Descripción General', style='TextoNormal')
doc.add_paragraph('    1.2 Propósito y Alcance', style='TextoNormal')
doc.add_paragraph('    1.3 Objetivos del Sistema', style='TextoNormal')

doc.add_paragraph('2. DOCUMENTACIÓN FUNCIONAL', style='Titulo1')
doc.add_paragraph('    2.1 Módulo de Usuarios', style='TextoNormal')
doc.add_paragraph('    2.2 Módulo de Novelas', style='TextoNormal')
doc.add_paragraph('    2.3 Módulo de Capítulos', style='TextoNormal')
doc.add_paragraph('    2.4 Módulo de Géneros', style='TextoNormal')
doc.add_paragraph('    2.5 Módulo de Interacción', style='TextoNormal')
doc.add_paragraph('    2.6 Módulo de Imágenes', style='TextoNormal')
doc.add_paragraph('    2.7 Perfiles de Usuario', style='TextoNormal')
doc.add_paragraph('    2.8 Flujos de Trabajo', style='TextoNormal')

doc.add_paragraph('3. DOCUMENTACIÓN TÉCNICA', style='Titulo1')
doc.add_paragraph('    3.1 Arquitectura del Sistema', style='TextoNormal')
doc.add_paragraph('    3.2 Stack Tecnológico', style='TextoNormal')
doc.add_paragraph('    3.3 Modelos y Base de Datos', style='TextoNormal')
doc.add_paragraph('    3.4 Estructura del Código', style='TextoNormal')
doc.add_paragraph('    3.5 API y Puntos de Entrada', style='TextoNormal')
doc.add_paragraph('    3.6 Seguridad y Autenticación', style='TextoNormal')
doc.add_paragraph('    3.7 Implementación de Características', style='TextoNormal')

doc.add_paragraph('4. ANEXOS', style='Titulo1')
doc.add_paragraph('    4.1 Glosario de Términos', style='TextoNormal')
doc.add_paragraph('    4.2 Diagramas', style='TextoNormal')

# Salto de página después del índice
doc.add_page_break()

# ----- INTRODUCCIÓN -----
doc.add_paragraph('1. INTRODUCCIÓN', style='TituloPrincipal')

doc.add_paragraph('1.1 Descripción General', style='Titulo2')
p = doc.add_paragraph(
    'NovelasApp es una plataforma web especializada para la creación, publicación y lectura de novelas en línea. '
    'Desarrollada con Laravel 10, el sistema permite a los escritores compartir sus obras por capítulos y a los '
    'lectores descubrir y seguir historias según sus géneros preferidos.',
    style='TextoNormal'
)

doc.add_paragraph('1.2 Propósito y Alcance', style='Titulo2')
p = doc.add_paragraph('El propósito principal de NovelasApp es proporcionar un espacio donde:', style='TextoNormal')
doc.add_paragraph('• Los escritores puedan publicar sus novelas capítulo a capítulo', style='TextoNormal')
doc.add_paragraph('• Los lectores puedan descubrir y disfrutar de nuevas historias', style='TextoNormal')
doc.add_paragraph('• Se fomente la comunidad literaria mediante interacciones básicas', style='TextoNormal')
doc.add_paragraph('• Se mantenga un catálogo organizado por géneros literarios', style='TextoNormal')

doc.add_paragraph('1.3 Objetivos del Sistema', style='Titulo2')
doc.add_paragraph('• Facilitar la creación y publicación de contenido literario', style='TextoNormal')
doc.add_paragraph('• Ofrecer una experiencia de lectura limpia y agradable', style='TextoNormal')
doc.add_paragraph('• Proporcionar herramientas para organizar y gestionar novelas', style='TextoNormal')
doc.add_paragraph('• Permitir la interacción básica entre usuarios y contenidos', style='TextoNormal')
doc.add_paragraph('• Establecer un sistema de categorización eficiente', style='TextoNormal')

# Salto de página
doc.add_page_break()

# ----- DOCUMENTACIÓN FUNCIONAL -----
doc.add_paragraph('2. DOCUMENTACIÓN FUNCIONAL', style='TituloPrincipal')

doc.add_paragraph('2.1 Módulo de Usuarios', style='Titulo2')
doc.add_paragraph('2.1.1 Registro e Inicio de Sesión', style='Titulo3')
doc.add_paragraph(
    'El sistema utiliza Laravel Breeze para proporcionar un sistema de autenticación completo y seguro. Las '
    'funcionalidades principales incluyen:',
    style='TextoNormal'
)
doc.add_paragraph('• Registro de usuario con nombre, correo electrónico y contraseña', style='TextoNormal')
doc.add_paragraph('• Inicio de sesión con credenciales verificadas', style='TextoNormal')
doc.add_paragraph('• Recuperación de contraseña mediante correo electrónico', style='TextoNormal')
doc.add_paragraph('• Opción para mantener la sesión activa entre visitas ("Recordarme")', style='TextoNormal')

doc.add_paragraph('2.1.2 Gestión de Perfil', style='Titulo3')
doc.add_paragraph('El usuario registrado puede gestionar su información personal a través de las siguientes opciones:', style='TextoNormal')
doc.add_paragraph('• Ver perfil con datos personales y estadísticas', style='TextoNormal')
doc.add_paragraph('• Editar información personal', style='TextoNormal')
doc.add_paragraph('• Cambiar contraseña', style='TextoNormal')
doc.add_paragraph('• Eliminar cuenta (dar de baja del sistema)', style='TextoNormal')

doc.add_paragraph('2.2 Módulo de Novelas', style='Titulo2')
doc.add_paragraph('2.2.1 Gestión de Novelas', style='Titulo3')
doc.add_paragraph(
    'Los usuarios registrados pueden crear y administrar sus propias novelas a través de un conjunto completo de '
    'herramientas CRUD (Crear, Leer, Actualizar, Eliminar):',
    style='TextoNormal'
)
doc.add_paragraph('• Creación de novela con título, descripción, portada y géneros', style='TextoNormal')
doc.add_paragraph('• Listado de novelas personales', style='TextoNormal')
doc.add_paragraph('• Edición de datos de la novela', style='TextoNormal')
doc.add_paragraph('• Eliminación de novelas completas', style='TextoNormal')
doc.add_paragraph('• Visualización de estadísticas básicas (lecturas, interacciones)', style='TextoNormal')

doc.add_paragraph('2.2.2 Exploración de Novelas', style='Titulo3')
doc.add_paragraph(
    'El sistema proporciona diferentes formas de descubrir el contenido disponible:',
    style='TextoNormal'
)
doc.add_paragraph('• Catálogo general de novelas', style='TextoNormal')
doc.add_paragraph('• Filtrado por género literario', style='TextoNormal')
doc.add_paragraph('• Búsqueda por texto', style='TextoNormal')
doc.add_paragraph('• Sección de novelas destacadas', style='TextoNormal')

# Continuación con más módulos funcionales...
doc.add_paragraph('2.3 Módulo de Capítulos', style='Titulo2')
doc.add_paragraph('2.3.1 Gestión de Capítulos', style='Titulo3')
doc.add_paragraph(
    'Cada novela puede contener múltiples capítulos que son gestionados por su autor:',
    style='TextoNormal'
)
doc.add_paragraph('• Creación de capítulos con editor de texto enriquecido', style='TextoNormal')
doc.add_paragraph('• Listado de capítulos por novela', style='TextoNormal')
doc.add_paragraph('• Edición de capítulos existentes', style='TextoNormal')
doc.add_paragraph('• Eliminación de capítulos', style='TextoNormal')
doc.add_paragraph('• Sistema de ordenación automática por número de capítulo', style='TextoNormal')

# ----- DOCUMENTACIÓN TÉCNICA -----
doc.add_page_break()
doc.add_paragraph('3. DOCUMENTACIÓN TÉCNICA', style='TituloPrincipal')

doc.add_paragraph('3.1 Arquitectura del Sistema', style='Titulo2')
doc.add_paragraph(
    'NovelasApp sigue una arquitectura Modelo-Vista-Controlador (MVC) implementada a través del framework Laravel 10. '
    'Esta arquitectura separa la lógica de negocio de la interfaz de usuario, facilitando el desarrollo y mantenimiento.',
    style='TextoNormal'
)

doc.add_paragraph('3.1.1 Componentes Principales', style='Titulo3')
doc.add_paragraph('• Capa de Presentación: Blade templates y TailwindCSS', style='TextoNormal')
doc.add_paragraph('• Capa de Negocio: Controladores organizados por funcionalidad', style='TextoNormal')
doc.add_paragraph('• Capa de Datos: Modelos Eloquent con relaciones definidas', style='TextoNormal')
doc.add_paragraph('• Capa de Persistencia: Base de datos MySQL', style='TextoNormal')

doc.add_paragraph('3.2 Stack Tecnológico', style='Titulo2')
doc.add_paragraph('El sistema está construido utilizando las siguientes tecnologías:', style='TextoNormal')
doc.add_paragraph('• Backend: PHP 8.1, Laravel 10', style='TextoNormal')
doc.add_paragraph('• Frontend: HTML5, CSS3, JavaScript, TailwindCSS, Alpine.js', style='TextoNormal')
doc.add_paragraph('• Base de Datos: MySQL', style='TextoNormal')
doc.add_paragraph('• Autenticación: Laravel Breeze', style='TextoNormal')
doc.add_paragraph('• Herramientas Adicionales: TinyMCE (editor WYSIWYG)', style='TextoNormal')

doc.add_paragraph('3.3 Modelos y Base de Datos', style='Titulo2')
doc.add_paragraph('3.3.1 Entidades Principales', style='Titulo3')
doc.add_paragraph('• User: Almacena información de usuarios y credenciales', style='TextoNormal')
doc.add_paragraph('• Novela: Contiene información sobre las novelas publicadas', style='TextoNormal')
doc.add_paragraph('• Capitulo: Guarda el contenido de cada capítulo de las novelas', style='TextoNormal')
doc.add_paragraph('• Genero: Categorías para clasificar las novelas', style='TextoNormal')
doc.add_paragraph('• CapituloImagen: Imágenes asociadas a cada capítulo', style='TextoNormal')

doc.add_paragraph('3.3.2 Relaciones Principales', style='Titulo3')
doc.add_paragraph('• Usuario → Novelas: Un usuario puede tener muchas novelas (1:N)', style='TextoNormal')
doc.add_paragraph('• Novela → Capítulos: Una novela contiene muchos capítulos (1:N)', style='TextoNormal')
doc.add_paragraph('• Novela ↔ Géneros: Relación muchos a muchos (N:M)', style='TextoNormal')
doc.add_paragraph('• Capítulo → Imágenes: Un capítulo puede tener múltiples imágenes (1:N)', style='TextoNormal')

doc.add_paragraph('3.4 Estructura del Código', style='Titulo2')
code_text = '''
novelas-app/
├── app/
│   ├── Http/
│   │   ├── Controllers/      # Controladores de la aplicación
│   │   ├── Middleware/       # Middleware de autenticación y permisos
│   │   └── Requests/         # Validación de formularios
│   ├── Models/               # Modelos Eloquent
│   └── Providers/            # Proveedores de servicios
├── config/                   # Configuración de la aplicación
├── database/
│   ├── migrations/           # Migraciones de la base de datos
│   └── seeders/              # Datos iniciales
├── public/                   # Archivos accesibles públicamente
│   └── storage/              # Imágenes y archivos subidos
├── resources/
│   ├── css/                  # Estilos CSS/TailwindCSS
│   ├── js/                   # Código JavaScript
│   └── views/                # Vistas Blade
├── routes/                   # Definición de rutas
│   └── web.php               # Rutas principales
└── storage/                  # Almacenamiento de la aplicación
'''
p = doc.add_paragraph(code_text, style='TextoCodigo')

doc.add_paragraph('3.5 API y Puntos de Entrada', style='Titulo2')
doc.add_paragraph('3.5.1 Rutas Principales', style='Titulo3')
doc.add_paragraph('El sistema organiza las rutas en grupos según su accesibilidad y funcionalidad:', style='TextoNormal')
doc.add_paragraph('• Rutas Públicas: Accesibles sin autenticación', style='TextoNormal')
doc.add_paragraph('• Rutas Protegidas: Requieren autenticación', style='TextoNormal')
doc.add_paragraph('• Rutas Administrativas: Restringidas a administradores', style='TextoNormal')

p = doc.add_paragraph(
    'Es importante destacar que las rutas de capítulos están anidadas dentro de las rutas de novelas, reflejando la '
    'relación jerárquica entre estos recursos. Por ejemplo: ',
    style='TextoNormal'
)
doc.add_paragraph('/novelas/{novela}/capitulos', style='TextoCodigo')

doc.add_paragraph('3.5.2 Middlewares', style='Titulo3')
doc.add_paragraph('Los principales middleware utilizados incluyen:', style='TextoNormal')
doc.add_paragraph('• auth: Verifica la autenticación del usuario', style='TextoNormal')
doc.add_paragraph('• verified: Confirma que el email del usuario ha sido verificado', style='TextoNormal')
doc.add_paragraph('• admin: Restringe el acceso a usuarios con rol de administrador', style='TextoNormal')

doc.add_paragraph('3.6 Seguridad y Autenticación', style='Titulo2')
doc.add_paragraph(
    'El sistema implementa varias capas de seguridad para proteger los datos y el acceso:',
    style='TextoNormal'
)
doc.add_paragraph('• Autenticación mediante Laravel Breeze', style='TextoNormal')
doc.add_paragraph('• Protección CSRF en formularios', style='TextoNormal')
doc.add_paragraph('• Validación de datos de entrada', style='TextoNormal')
doc.add_paragraph('• Almacenamiento seguro de contraseñas con bcrypt', style='TextoNormal')
doc.add_paragraph('• Sesiones con tiempo de expiración configurable', style='TextoNormal')

# ANEXOS
doc.add_page_break()
doc.add_paragraph('4. ANEXOS', style='TituloPrincipal')

doc.add_paragraph('4.1 Glosario de Términos', style='Titulo2')
# Crear una tabla para el glosario
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Término'
hdr_cells[1].text = 'Definición'

# Agregar términos al glosario
terms = [
    ('Laravel', 'Framework de PHP utilizado para el desarrollo de la aplicación.'),
    ('MVC', 'Patrón de diseño Modelo-Vista-Controlador implementado en la arquitectura.'),
    ('Middleware', 'Mecanismo de filtrado de peticiones HTTP en Laravel.'),
    ('Eloquent', 'ORM (Object-Relational Mapping) utilizado para interactuar con la base de datos.'),
    ('Blade', 'Motor de plantillas de Laravel utilizado para las vistas.'),
    ('TailwindCSS', 'Framework CSS utilizado para el diseño de la interfaz de usuario.'),
    ('MySQL', 'Sistema de gestión de base de datos relacional utilizado para almacenar los datos.'),
    ('Breeze', 'Paquete de Laravel que proporciona implementación de autenticación.'),
]

for term, definition in terms:
    row_cells = table.add_row().cells
    row_cells[0].text = term
    row_cells[1].text = definition

# Pie de página
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = f'NovelasApp - Documentación Técnica y Funcional - Generada el {current_date}'
footer_para.style = doc.styles['PiePagina']
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Guardar el documento
doc_path = os.path.join(doc_dir, 'NovelasApp-Documentacion.docx')
doc.save(doc_path)
print(f"Documentación generada con éxito en: {doc_path}")
